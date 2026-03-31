"""Тесты для TemplateEngine — работа с DOCX-шаблонами."""

from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from src.template_engine import TemplateEngine


# ---------------------------------------------------------------------------
# Фикстуры — создание тестовых DOCX
# ---------------------------------------------------------------------------

@pytest.fixture()
def doc_with_paragraph(tmp_path: Path) -> Path:
    """DOCX с одним параграфом, содержащим плейсхолдер [NAME]."""
    path = tmp_path / "paragraph.docx"
    doc = Document()
    doc.add_paragraph("[NAME]")
    doc.save(str(path))
    return path


@pytest.fixture()
def doc_with_table(tmp_path: Path) -> Path:
    """DOCX с таблицей 2x2, одна ячейка содержит [VALUE]."""
    path = tmp_path / "table.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Заголовок"
    table.cell(0, 1).text = "[VALUE]"
    table.cell(1, 0).text = "Строка 2"
    table.cell(1, 1).text = "Данные"
    doc.save(str(path))
    return path


@pytest.fixture()
def doc_with_bold_run(tmp_path: Path) -> Path:
    """DOCX с жирным Run-ом, содержащим плейсхолдер [X]."""
    path = tmp_path / "bold.docx"
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("[X]")
    run.bold = True
    run.font.size = Pt(14)
    doc.save(str(path))
    return path


@pytest.fixture()
def doc_with_multiple_placeholders(tmp_path: Path) -> Path:
    """DOCX с одним плейсхолдером в 2 местах (параграф + таблица)."""
    path = tmp_path / "multiple.docx"
    doc = Document()
    doc.add_paragraph("Компания: [COMPANY]")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Компания: [COMPANY]"
    doc.save(str(path))
    return path


@pytest.fixture()
def doc_for_table_row(tmp_path: Path) -> Path:
    """DOCX с таблицей 2x3 для тестирования fill_table_row."""
    path = tmp_path / "table_row.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "№"
    table.cell(0, 1).text = "Наименование"
    table.cell(0, 2).text = "Кол-во"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(1, 2).text = ""
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Тесты
# ---------------------------------------------------------------------------

class TestReplacePlaceholder:
    """Тесты для TemplateEngine.replace_placeholder."""

    def test_replace_in_paragraph(self, doc_with_paragraph: Path) -> None:
        """Плейсхолдер в параграфе заменяется на значение."""
        engine = TemplateEngine(doc_with_paragraph)
        engine.replace_placeholder("[NAME]", "ООО «Тест»")

        # Проверяем через внутренний документ
        text = engine._doc.paragraphs[0].text
        assert text == "ООО «Тест»"
        assert "[NAME]" not in text

    def test_replace_in_table(self, doc_with_table: Path) -> None:
        """Плейсхолдер в ячейке таблицы заменяется на значение."""
        engine = TemplateEngine(doc_with_table)
        engine.replace_placeholder("[VALUE]", "12345")

        cell_text = engine._doc.tables[0].cell(0, 1).text
        assert cell_text == "12345"

    def test_replace_preserves_bold(self, doc_with_bold_run: Path) -> None:
        """После замены жирный шрифт и размер сохраняются."""
        engine = TemplateEngine(doc_with_bold_run)
        engine.replace_placeholder("[X]", "Значение")

        run = engine._doc.paragraphs[0].runs[0]
        assert run.text == "Значение"
        assert run.bold is True
        assert run.font.size == Pt(14)

    def test_replace_count(self, doc_with_multiple_placeholders: Path) -> None:
        """Метод возвращает количество произведённых замен."""
        engine = TemplateEngine(doc_with_multiple_placeholders)
        count = engine.replace_placeholder("[COMPANY]", "ООО «Ромашка»")

        assert count == 2


class TestFillTableRow:
    """Тесты для TemplateEngine.fill_table_row."""

    def test_fill_row_by_columns(self, doc_for_table_row: Path) -> None:
        """Заполнение ячеек строки таблицы по индексам колонок."""
        engine = TemplateEngine(doc_for_table_row)
        engine.fill_table_row(
            table_idx=0,
            row_idx=1,
            col_values={0: "1", 1: "Кабель ВВГнг", 2: "100"},
        )

        table = engine._doc.tables[0]
        assert table.cell(1, 0).text == "1"
        assert table.cell(1, 1).text == "Кабель ВВГнг"
        assert table.cell(1, 2).text == "100"

    def test_fill_row_partial(self, doc_for_table_row: Path) -> None:
        """Частичное заполнение — только указанные колонки."""
        engine = TemplateEngine(doc_for_table_row)
        engine.fill_table_row(
            table_idx=0,
            row_idx=1,
            col_values={1: "Кабель"},
        )

        table = engine._doc.tables[0]
        # Колонка 0 остаётся пустой
        assert table.cell(1, 0).text == ""
        assert table.cell(1, 1).text == "Кабель"

    def test_fill_row_invalid_table_index(self, doc_for_table_row: Path) -> None:
        """Невалидный индекс таблицы — IndexError."""
        engine = TemplateEngine(doc_for_table_row)
        with pytest.raises(IndexError):
            engine.fill_table_row(table_idx=99, row_idx=0, col_values={0: "X"})


class TestSave:
    """Тесты для TemplateEngine.save."""

    def test_save_creates_file(self, doc_with_paragraph: Path, tmp_path: Path) -> None:
        """Сохранение создаёт файл по указанному пути."""
        output = tmp_path / "sub" / "output.docx"
        engine = TemplateEngine(doc_with_paragraph)
        engine.save(output)

        assert output.exists()
        assert output.stat().st_size > 0

    def test_original_not_modified(
        self, doc_with_paragraph: Path, tmp_path: Path
    ) -> None:
        """После замены и сохранения оригинальный шаблон не изменён."""
        output = tmp_path / "output.docx"
        engine = TemplateEngine(doc_with_paragraph)
        engine.replace_placeholder("[NAME]", "Замена")
        engine.save(output)

        # Перечитываем оригинал — плейсхолдер должен быть на месте
        original = Document(str(doc_with_paragraph))
        assert "[NAME]" in original.paragraphs[0].text
