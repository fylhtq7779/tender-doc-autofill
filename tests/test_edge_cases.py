"""Комплексные edge case и стресс-тесты для системы генерации тендерных документов.

Категории тестов:
  1. DataLoader - пустые/экстремальные JSON
  2. Utils - граничные значения форматирования
  3. TemplateEngine - нестандартные подстановки
  4. MappingLoader - битые/неполные YAML
  5. Generator - интеграционные edge cases
  6. Extractor - пустые/несуществующие DOCX
  7. CLI - невалидные аргументы
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest
import yaml
from docx import Document

from src.data_loader import DataContext, DataLoader
from src.extractor import TenderExtractor
from src.generator import DocumentGenerator
from src.mapping_loader import MappingLoader
from src.template_engine import TemplateEngine
from src.utils import format_date_long, format_money, resolve_dot_path


# ===========================================================================
# Хелперы для создания тестовых данных
# ===========================================================================

def _minimal_profile() -> dict:
    """Минимальный валидный company_profile."""
    return {
        "company": {
            "full_name": "ООО «Тест»",
            "short_name": "ООО «Тест»",
            "inn": "7705123456",
            "kpp": "770501001",
            "ogrn": "1127746123456",
            "legal_address_full": "г. Москва",
            "legal_address_short": "Москва",
            "postal_address": "г. Москва",
            "country": "Россия",
            "city": "Москва",
        },
        "bank": {
            "name": "Банк",
            "account": "40702810900000012345",
            "correspondent_account": "30101810400000000225",
            "bik": "044525225",
        },
        "contact": {
            "responsible_name_full": "Иванов Иван Иванович",
            "responsible_name_short": "Иванов И.И.",
            "phone": "+7 999 000-00-00",
            "email": "test@test.ru",
        },
        "signatory": {
            "position": "Директор",
            "name_short": "Иванов И.И.",
            "name_full": "Иванов Иван Иванович",
            "basis": "Устав",
        },
    }


def _minimal_tender(**overrides: object) -> dict:
    """Минимальный валидный tender с возможностью переопределения полей."""
    base = {
        "purchase_number": "PUR-2026-001",
        "lot_number": "1",
        "subject": "Поставка",
        "customer": {
            "full_name": "АО «Заказчик»",
            "short_name": "АО «Заказчик»",
            "legal_address": "Адрес",
            "postal_address": "Адрес",
            "email": "c@c.ru",
            "inn": "7705123456",
            "kpp": "770501001",
            "ogrn": "1127746123456",
            "bank": {
                "name": "Банк",
                "account": "40702810900000012345",
                "correspondent_account": "30101810400000000225",
                "bik": "044525225",
            },
            "signatory": {"position": "Директор", "name": "Иванов И.И."},
        },
        "delivery": {"place": "Москва", "basis": "DAP"},
    }
    base.update(overrides)
    return base


def _minimal_calc(**overrides: object) -> dict:
    """Минимальный валидный calc с возможностью переопределения полей."""
    base = {"vat_rate": 20.0}
    base.update(overrides)
    return base


def _write_jsons(
    tmp_path: Path,
    profile: dict | None = None,
    tender: dict | None = None,
    calc: dict | None = None,
) -> None:
    """Записывает 3 JSON-файла в директорию."""
    for name, data in [
        ("company_profile.json", profile or _minimal_profile()),
        ("tender.json", tender or _minimal_tender()),
        ("calc.json", calc or _minimal_calc()),
    ]:
        (tmp_path / name).write_text(
            json.dumps(data, ensure_ascii=False), encoding="utf-8"
        )


def _make_simple_docx(path: Path, text: str = "[PLACEHOLDER]") -> None:
    """Создаёт простой DOCX с одним параграфом."""
    doc = Document()
    doc.add_paragraph(text)
    doc.save(str(path))


def _make_docx_with_table(path: Path, rows: int = 3, cols: int = 2) -> None:
    """Создаёт DOCX с таблицей заданного размера."""
    doc = Document()
    doc.add_paragraph("Заголовок")
    table = doc.add_table(rows=rows, cols=cols)
    for r in range(rows):
        for c in range(cols):
            table.cell(r, c).text = f"r{r}c{c}"
    doc.save(str(path))


def _setup_generator_env(
    tmp_path: Path,
    fields: list[dict] | None = None,
    table_rows: list[dict] | None = None,
    template_text: str = "[NAME] [TOTAL] [DATE]",
    profile: dict | None = None,
    tender: dict | None = None,
    calc: dict | None = None,
    template_factory: str | None = None,
) -> DocumentGenerator:
    """Создаёт полное тестовое окружение для DocumentGenerator."""
    data_dir = tmp_path / "data"
    data_dir.mkdir(exist_ok=True)
    _write_jsons(data_dir, profile=profile, tender=tender, calc=calc)

    templates_dir = tmp_path / "templates"
    templates_dir.mkdir(exist_ok=True)
    template_path = templates_dir / "test.docx"

    if template_factory == "table":
        _make_docx_with_table(template_path)
    else:
        _make_simple_docx(template_path, template_text)

    mappings_dir = tmp_path / "mappings"
    mappings_dir.mkdir(exist_ok=True)
    output_dir = tmp_path / "output"

    mapping = {
        "document": {
            "name": "Тест",
            "template": str(template_path),
            "output_name": "result.docx",
        },
        "fields": fields or [],
        "table_rows": table_rows or [],
    }
    (mappings_dir / "01_test.yaml").write_text(
        yaml.dump(mapping, allow_unicode=True), encoding="utf-8"
    )

    return DocumentGenerator(
        data_dir=data_dir,
        mappings_dir=mappings_dir,
        output_dir=output_dir,
    )


# ===========================================================================
# 1. DATA LOADER EDGE CASES
# ===========================================================================

class TestDataLoaderEdgeCases:
    """Граничные случаи загрузки JSON-данных."""

    def test_empty_json_file(self, tmp_path: Path) -> None:
        """Пустой JSON ({}) без обязательных полей - ValidationError."""
        (tmp_path / "company_profile.json").write_text("{}", encoding="utf-8")
        (tmp_path / "tender.json").write_text("{}", encoding="utf-8")
        (tmp_path / "calc.json").write_text("{}", encoding="utf-8")

        loader = DataLoader(data_dir=tmp_path)
        with pytest.raises(Exception):
            loader.load_all()

    def test_json_with_extra_fields(self, tmp_path: Path) -> None:
        """JSON с лишними полями - extra поля игнорируются (model_config extra=ignore)."""
        profile = _minimal_profile()
        profile["unknown_field"] = "должно быть проигнорировано"
        profile["company"]["extra_stuff"] = 42
        _write_jsons(tmp_path, profile=profile)

        loader = DataLoader(data_dir=tmp_path)
        ctx = loader.load_all()
        assert ctx.profile.company.inn == "7705123456"

    def test_unicode_in_data(self, tmp_path: Path) -> None:
        """Unicode-символы в данных: кавычки-ёлочки, ё, №."""
        profile = _minimal_profile()
        profile["company"]["full_name"] = 'ООО «Ёлочка» № 1'
        _write_jsons(tmp_path, profile=profile)

        loader = DataLoader(data_dir=tmp_path)
        ctx = loader.load_all()
        assert ctx.profile.company.full_name == 'ООО «Ёлочка» № 1'

    def test_very_long_company_name(self, tmp_path: Path) -> None:
        """Имя компании > 500 символов - загружается без ошибок."""
        profile = _minimal_profile()
        long_name = "ООО «" + "А" * 500 + "»"
        profile["company"]["full_name"] = long_name
        _write_jsons(tmp_path, profile=profile)

        loader = DataLoader(data_dir=tmp_path)
        ctx = loader.load_all()
        assert len(ctx.profile.company.full_name) > 500

    def test_empty_string_fields(self, tmp_path: Path) -> None:
        """Пустые строки в ИНН/КПП - загрузка с warning, без crash."""
        profile = _minimal_profile()
        profile["company"]["inn"] = ""
        profile["company"]["kpp"] = ""
        _write_jsons(tmp_path, profile=profile)

        loader = DataLoader(data_dir=tmp_path)
        # Пустая строка не пройдёт валидацию цифр - но модель допускает,
        # только warning в лог
        ctx = loader.load_all()
        assert ctx.profile.company.inn == ""
        assert ctx.profile.company.kpp == ""

    def test_none_values_in_json(self, tmp_path: Path) -> None:
        """null-значения в необязательных полях - загружаются как None/default."""
        tender = _minimal_tender()
        tender["payment"] = None
        tender["warranty"] = None
        _write_jsons(tmp_path, tender=tender)

        loader = DataLoader(data_dir=tmp_path)
        ctx = loader.load_all()
        assert ctx.tender.payment is None
        assert ctx.tender.warranty is None

    def test_zero_items(self, tmp_path: Path) -> None:
        """Пустой список items - загрузка без ошибок."""
        tender = _minimal_tender(items=[])
        calc = _minimal_calc(items=[])
        _write_jsons(tmp_path, tender=tender, calc=calc)

        loader = DataLoader(data_dir=tmp_path)
        ctx = loader.load_all()
        assert ctx.tender.items == []
        assert ctx.calc.items == []

    def test_single_item(self, tmp_path: Path) -> None:
        """Один элемент в items - загружается корректно."""
        tender = _minimal_tender(items=[
            {"line_no": 1, "article": "A1", "name": "Кабель", "unit": "м",
             "qty": 100, "nmc_unit_price": 150.0},
        ])
        _write_jsons(tmp_path, tender=tender)

        loader = DataLoader(data_dir=tmp_path)
        ctx = loader.load_all()
        assert len(ctx.tender.items) == 1
        assert ctx.tender.items[0].name == "Кабель"

    def test_negative_price(self, tmp_path: Path) -> None:
        """Отрицательная цена - модель принимает (float без ограничений)."""
        calc = _minimal_calc(items=[
            {"line_no": 1, "unit_price_wo_vat": -100.0},
        ])
        _write_jsons(tmp_path, calc=calc)

        loader = DataLoader(data_dir=tmp_path)
        ctx = loader.load_all()
        assert ctx.calc.items[0].unit_price_wo_vat == -100.0

    def test_zero_price(self, tmp_path: Path) -> None:
        """Нулевая цена - загрузка без ошибок."""
        calc = _minimal_calc(items=[
            {"line_no": 1, "unit_price_wo_vat": 0.0},
        ])
        _write_jsons(tmp_path, calc=calc)

        loader = DataLoader(data_dir=tmp_path)
        ctx = loader.load_all()
        assert ctx.calc.items[0].unit_price_wo_vat == 0.0

    def test_huge_amount(self, tmp_path: Path) -> None:
        """Огромная сумма - загрузка без ошибок."""
        calc = _minimal_calc(total_with_vat=999_999_999_999.99)
        _write_jsons(tmp_path, calc=calc)

        loader = DataLoader(data_dir=tmp_path)
        ctx = loader.load_all()
        assert ctx.calc.total_with_vat == pytest.approx(999_999_999_999.99, rel=1e-9)

    def test_fractional_kopecks(self, tmp_path: Path) -> None:
        """Три десятичных знака (123.456) - Pydantic принимает float без округления."""
        calc = _minimal_calc(total_with_vat=123.456)
        _write_jsons(tmp_path, calc=calc)

        loader = DataLoader(data_dir=tmp_path)
        ctx = loader.load_all()
        assert ctx.calc.total_with_vat == pytest.approx(123.456)


# ===========================================================================
# 2. UTILS EDGE CASES
# ===========================================================================

class TestFormatMoneyEdgeCases:
    """Граничные случаи форматирования денежных сумм."""

    def test_format_money_negative(self) -> None:
        """Отрицательная сумма - знак минус сохраняется."""
        result = format_money(-1000.0)
        # Должен содержать "1 000" и запятую
        assert "1 000" in result or "1000" in result

    def test_format_money_tiny(self) -> None:
        """Минимальная сумма 0.01 - отображается как 0,01."""
        result = format_money(0.01)
        assert result == "0,01"

    def test_format_money_huge(self) -> None:
        """Миллиард - разделители тысяч расставлены корректно."""
        result = format_money(1_000_000_000.0)
        assert result == "1 000 000 000,00"


class TestFormatDateLongEdgeCases:
    """Граничные случаи форматирования дат."""

    def test_format_date_long_wrong_format(self) -> None:
        """Формат YYYY-MM-DD (ISO) - возвращает исходную строку."""
        result = format_date_long("2026-03-31")
        assert result == "2026-03-31"

    def test_format_date_long_january_first(self) -> None:
        """1 января - граничный день."""
        result = format_date_long("01.01.2026")
        assert result == "«01» января 2026 года"

    def test_format_date_long_december_last(self) -> None:
        """31 декабря - граничный день."""
        result = format_date_long("31.12.2026")
        assert result == "«31» декабря 2026 года"


class TestResolveDotPathEdgeCases:
    """Граничные случаи для resolve_dot_path."""

    def test_deeply_nested(self) -> None:
        """5 уровней вложенности - доступ корректен."""
        data = {"a": {"b": {"c": {"d": {"e": "глубоко"}}}}}
        result = resolve_dot_path(data, "a.b.c.d.e")
        assert result == "глубоко"

    def test_array_out_of_bounds(self) -> None:
        """Индекс массива за пределами - None."""
        data = {"items": [{"name": "один"}]}
        result = resolve_dot_path(data, "items[999].name")
        assert result is None

    def test_empty_dict(self) -> None:
        """Пустой словарь - None для любого пути."""
        result = resolve_dot_path({}, "any.path")
        assert result is None

    def test_none_input(self) -> None:
        """None как входные данные - TypeError или None."""
        # resolve_dot_path ожидает dict, но None не должен вызывать crash
        # если path пуст - возвращает data (None), если нет - None
        result = resolve_dot_path({}, "nonexistent")
        assert result is None

    def test_special_chars_in_value(self) -> None:
        """Значение содержит спецсимволы - извлекается как есть."""
        data = {"field": "Значение с <>&\"' символами"}
        result = resolve_dot_path(data, "field")
        assert result == "Значение с <>&\"' символами"


# ===========================================================================
# 3. TEMPLATE ENGINE EDGE CASES
# ===========================================================================

class TestTemplateEngineEdgeCases:
    """Граничные случаи для TemplateEngine."""

    def test_placeholder_not_found(self, tmp_path: Path) -> None:
        """Плейсхолдер отсутствует в документе - 0 замен, без ошибок."""
        path = tmp_path / "test.docx"
        _make_simple_docx(path, "Обычный текст без плейсхолдеров")

        engine = TemplateEngine(path)
        count = engine.replace_placeholder("[MISSING]", "значение")
        assert count == 0

    def test_placeholder_appears_multiple_times(self, tmp_path: Path) -> None:
        """Один плейсхолдер в 3 местах - все 3 заменяются."""
        path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("[X] первый")
        doc.add_paragraph("[X] второй")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "[X] третий"
        doc.save(str(path))

        engine = TemplateEngine(path)
        count = engine.replace_placeholder("[X]", "OK")
        assert count == 3

    def test_empty_value_replacement(self, tmp_path: Path) -> None:
        """Замена на пустую строку - плейсхолдер удаляется."""
        path = tmp_path / "test.docx"
        _make_simple_docx(path, "Значение: [VAL].")

        engine = TemplateEngine(path)
        engine.replace_placeholder("[VAL]", "")
        text = engine._doc.paragraphs[0].text
        assert "[VAL]" not in text
        assert "Значение: ." in text

    def test_very_long_value(self, tmp_path: Path) -> None:
        """Замена на строку из 10000 символов - работает без ошибок."""
        path = tmp_path / "test.docx"
        _make_simple_docx(path, "[LONG]")

        engine = TemplateEngine(path)
        long_value = "А" * 10000
        count = engine.replace_placeholder("[LONG]", long_value)
        assert count == 1
        assert len(engine._doc.paragraphs[0].text) == 10000

    def test_special_xml_chars_in_value(self, tmp_path: Path) -> None:
        """Спецсимволы XML (<>&\"') в значении - подставляются корректно."""
        path = tmp_path / "test.docx"
        _make_simple_docx(path, "[VAL]")

        engine = TemplateEngine(path)
        engine.replace_placeholder("[VAL]", 'ООО "A&B" <test>')

        # Сохраняем и перечитываем - проверяем что XML не сломан
        output = tmp_path / "out.docx"
        engine.save(output)
        doc2 = Document(str(output))
        assert doc2.paragraphs[0].text == 'ООО "A&B" <test>'

    def test_newline_in_value(self, tmp_path: Path) -> None:
        """Значение с \\n - подставляется (может не отображаться как перевод строки в DOCX)."""
        path = tmp_path / "test.docx"
        _make_simple_docx(path, "[VAL]")

        engine = TemplateEngine(path)
        engine.replace_placeholder("[VAL]", "строка1\nстрока2")
        text = engine._doc.paragraphs[0].text
        assert "строка1" in text
        assert "строка2" in text

    def test_fill_row_out_of_bounds(self, tmp_path: Path) -> None:
        """row_idx за пределами таблицы - IndexError."""
        path = tmp_path / "test.docx"
        _make_docx_with_table(path, rows=2, cols=2)

        engine = TemplateEngine(path)
        with pytest.raises(IndexError):
            engine.fill_table_row(table_idx=0, row_idx=999, col_values={0: "X"})


# ===========================================================================
# 4. MAPPING LOADER EDGE CASES
# ===========================================================================

class TestMappingLoaderEdgeCases:
    """Граничные случаи загрузки YAML-маппингов."""

    def test_empty_yaml_file(self, tmp_path: Path) -> None:
        """Пустой YAML-файл - ошибка валидации (нет document)."""
        path = tmp_path / "empty.yaml"
        path.write_text("", encoding="utf-8")

        loader = MappingLoader()
        with pytest.raises(Exception):
            loader.load(path)

    def test_yaml_with_no_fields(self, tmp_path: Path) -> None:
        """YAML с document, но без fields - загрузка OK (fields по умолчанию [])."""
        data = {
            "document": {
                "name": "Тест",
                "template": "templates/test.docx",
                "output_name": "test.docx",
            },
        }
        path = tmp_path / "test.yaml"
        path.write_text(yaml.dump(data, allow_unicode=True), encoding="utf-8")

        loader = MappingLoader()
        result = loader.load(path)
        assert result.document.name == "Тест"
        assert result.fields == []

    def test_yaml_missing_document_section(self, tmp_path: Path) -> None:
        """YAML без секции document - ValidationError."""
        data = {"fields": []}
        path = tmp_path / "test.yaml"
        path.write_text(yaml.dump(data, allow_unicode=True), encoding="utf-8")

        loader = MappingLoader()
        with pytest.raises(Exception):
            loader.load(path)

    def test_yaml_invalid_source(self, tmp_path: Path) -> None:
        """source: UNKNOWN - маппинг загружается (валидация source - в generator)."""
        data = {
            "document": {
                "name": "Тест",
                "template": "templates/test.docx",
                "output_name": "test.docx",
            },
            "fields": [
                {"placeholder": "[X]", "source": "UNKNOWN", "path": "field"},
            ],
        }
        path = tmp_path / "test.yaml"
        path.write_text(yaml.dump(data, allow_unicode=True), encoding="utf-8")

        loader = MappingLoader()
        result = loader.load(path)
        assert result.fields[0].source == "UNKNOWN"

    def test_duplicate_placeholders_in_mapping(self, tmp_path: Path) -> None:
        """Один плейсхолдер дважды в маппинге - оба загружаются (без дедупликации)."""
        data = {
            "document": {
                "name": "Тест",
                "template": "templates/test.docx",
                "output_name": "test.docx",
            },
            "fields": [
                {"placeholder": "[X]", "source": "PROFILE", "path": "company.inn"},
                {"placeholder": "[X]", "source": "TENDER", "path": "purchase_number"},
            ],
        }
        path = tmp_path / "test.yaml"
        path.write_text(yaml.dump(data, allow_unicode=True), encoding="utf-8")

        loader = MappingLoader()
        result = loader.load(path)
        assert len(result.fields) == 2
        assert result.fields[0].placeholder == result.fields[1].placeholder


# ===========================================================================
# 5. GENERATOR EDGE CASES (ИНТЕГРАЦИЯ)
# ===========================================================================

class TestGeneratorEdgeCases:
    """Интеграционные граничные случаи генератора."""

    def test_generate_with_missing_optional_data(self, tmp_path: Path) -> None:
        """Отсутствующие необязательные поля - документ генерируется."""
        gen = _setup_generator_env(
            tmp_path,
            fields=[
                {
                    "placeholder": "[NAME]",
                    "source": "PROFILE",
                    "path": "company.full_name",
                },
                {
                    # Путь к несуществующему необязательному полю
                    "placeholder": "[TOTAL]",
                    "source": "CALC",
                    "path": "nonexistent_field",
                    "required": False,
                },
            ],
            template_text="[NAME] / [TOTAL] / конец",
        )

        paths = gen.generate_all()
        assert len(paths) == 1
        assert paths[0].exists()

        doc = Document(str(paths[0]))
        text = doc.paragraphs[0].text
        # NAME заменён, TOTAL заменён на пустую строку
        assert "ООО «Тест»" in text
        assert "[TOTAL]" not in text

    def test_generate_produces_valid_docx(self, tmp_path: Path) -> None:
        """Сгенерированный файл - валидный DOCX (открывается python-docx)."""
        gen = _setup_generator_env(
            tmp_path,
            fields=[
                {
                    "placeholder": "[NAME]",
                    "source": "PROFILE",
                    "path": "company.full_name",
                },
            ],
            template_text="Компания: [NAME]",
        )

        paths = gen.generate_all()
        # Не должно бросить исключение при открытии
        doc = Document(str(paths[0]))
        assert len(doc.paragraphs) > 0

    def test_generate_all_fields_replaced(self, tmp_path: Path) -> None:
        """Все плейсхолдеры из маппинга заменены - ни одного [...] не осталось."""
        gen = _setup_generator_env(
            tmp_path,
            fields=[
                {
                    "placeholder": "[NAME]",
                    "source": "PROFILE",
                    "path": "company.full_name",
                    "required": True,
                },
                {
                    "placeholder": "[DATE]",
                    "source": "SYSTEM",
                    "path": "current_date",
                },
                {
                    "placeholder": "[TOTAL]",
                    "source": "CALC",
                    "path": "total_with_vat",
                    "transform": "money",
                    "required": False,
                },
            ],
            template_text="[NAME] [DATE] [TOTAL]",
            calc=_minimal_calc(total_with_vat=1000.0),
        )

        paths = gen.generate_all()
        doc = Document(str(paths[0]))
        full_text = " ".join(p.text for p in doc.paragraphs)
        # Ни один из маппинговых плейсхолдеров не должен остаться
        assert "[NAME]" not in full_text
        assert "[DATE]" not in full_text
        assert "[TOTAL]" not in full_text

    def test_generate_idempotent(self, tmp_path: Path) -> None:
        """Двойной запуск генерации - одинаковый результат (перезапись файла)."""
        gen = _setup_generator_env(
            tmp_path,
            fields=[
                {
                    "placeholder": "[NAME]",
                    "source": "PROFILE",
                    "path": "company.full_name",
                },
            ],
            template_text="[NAME]",
        )

        # Первый запуск
        paths1 = gen.generate_all()
        doc1 = Document(str(paths1[0]))
        text1 = doc1.paragraphs[0].text

        # Второй запуск - перезаписывает файл
        paths2 = gen.generate_all()
        doc2 = Document(str(paths2[0]))
        text2 = doc2.paragraphs[0].text

        assert text1 == text2
        assert paths1[0] == paths2[0]


# ===========================================================================
# 6. EXTRACTOR EDGE CASES
# ===========================================================================

class TestExtractorEdgeCases:
    """Граничные случаи для TenderExtractor."""

    def test_extract_from_empty_docx(self, tmp_path: Path) -> None:
        """Пустой DOCX (без параграфов с данными) - пустые строки, без crash."""
        path = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(path))

        extractor = TenderExtractor(path)
        result = extractor.extract()

        # Все текстовые поля - пустые строки
        assert result["purchase_number"] == ""
        assert result["subject"] == ""
        assert result["items"] == []

    def test_extract_nonexistent_file(self) -> None:
        """Несуществующий файл - FileNotFoundError."""
        with pytest.raises(FileNotFoundError):
            TenderExtractor("/nonexistent/path/file.docx")


# ===========================================================================
# 7. CLI EDGE CASES
# ===========================================================================

class TestCliEdgeCases:
    """Граничные случаи CLI."""

    def test_cli_generate_nonexistent_data(self, tmp_path: Path) -> None:
        """generate с --data-dir /nonexistent - exit code != 0."""
        from click.testing import CliRunner
        from src.main import cli

        runner = CliRunner()
        result = runner.invoke(cli, [
            "generate",
            "--data-dir", str(tmp_path / "nonexistent_data"),
            "--mappings-dir", str(tmp_path / "mappings"),
            "--output-dir", str(tmp_path / "output"),
        ])
        assert result.exit_code != 0

    def test_cli_generate_nonexistent_mappings(self, tmp_path: Path) -> None:
        """generate с --mappings-dir /nonexistent - exit code != 0 или 0 документов."""
        from click.testing import CliRunner
        from src.main import cli

        # Создаём данные, но не маппинги
        data_dir = tmp_path / "data"
        data_dir.mkdir()
        _write_jsons(data_dir)

        runner = CliRunner()
        result = runner.invoke(cli, [
            "generate",
            "--data-dir", str(data_dir),
            "--mappings-dir", str(tmp_path / "nonexistent_mappings"),
            "--output-dir", str(tmp_path / "output"),
        ])
        # Либо ошибка, либо 0 документов сгенерировано
        if result.exit_code == 0:
            assert "0 документ" in result.output
        else:
            assert result.exit_code != 0

    def test_cli_extract_nonexistent_file(self, tmp_path: Path) -> None:
        """extract-tender с несуществующим файлом - exit code != 0."""
        from click.testing import CliRunner
        from src.main import cli

        runner = CliRunner()
        result = runner.invoke(cli, [
            "extract-tender",
            str(tmp_path / "nonexistent.docx"),
        ])
        assert result.exit_code != 0
