"""Тесты для DocumentGenerator - генерация документов по маппингу."""

import json
from pathlib import Path

import pytest
import yaml
from docx import Document

from src.generator import DocumentGenerator


# ---------------------------------------------------------------------------
# Фикстуры - создание тестового окружения
# ---------------------------------------------------------------------------

def _create_test_docx(path: Path) -> None:
    """Создаёт тестовый DOCX с плейсхолдерами и таблицей."""
    doc = Document()
    doc.add_paragraph("Компания: [NAME]")
    doc.add_paragraph("Сумма: [TOTAL]")
    doc.add_paragraph("Дата: [DATE]")
    # Таблица: заголовок + 2 строки данных
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "№"
    table.cell(0, 1).text = "Наименование"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = ""
    table.cell(2, 1).text = ""
    doc.save(str(path))


def _create_test_jsons(data_dir: Path) -> None:
    """Создаёт минимальные JSON-файлы для тестов."""
    profile = {
        "company": {
            "full_name": "ООО «Тестовая Компания»",
            "short_name": "ООО «ТК»",
            "inn": "7705123456",
            "kpp": "770501001",
            "ogrn": "1127746123456",
            "legal_address_full": "Адрес полный",
            "legal_address_short": "Адрес",
            "postal_address": "Почтовый адрес",
            "country": "Россия",
            "city": "Москва",
        },
        "bank": {
            "name": "АО «Тест-Банк»",
            "account": "40702810900000012345",
            "correspondent_account": "30101810400000000225",
            "bik": "044525225",
        },
        "contact": {
            "responsible_name_full": "Петров Иван Иванович",
            "responsible_name_short": "Петров И.И.",
            "phone": "+7 999 000-00-00",
            "email": "test@test.ru",
        },
        "signatory": {
            "position": "Генеральный директор",
            "name_short": "Петров И.И.",
            "name_full": "Петров Иван Иванович",
            "basis": "Устав",
        },
    }
    tender = {
        "purchase_number": "PUR-2026-001",
        "lot_number": "1",
        "subject": "Поставка кабеля",
        "customer": {
            "full_name": "АО «Заказчик»",
            "short_name": "АО «Заказчик»",
            "legal_address": "Адрес",
            "postal_address": "Адрес",
            "email": "c@c.ru",
            "inn": "7705123456",
            "kpp": "770501001",
            "ogrn": "1127746123456",
            "bank": {
                "name": "Банк",
                "account": "40702810900000012345",
                "correspondent_account": "30101810400000000225",
                "bik": "044525225",
            },
            "signatory": {"position": "Директор", "name": "Иванов И.И."},
        },
        "delivery": {"place": "Москва", "basis": "DAP"},
        "items": [
            {
                "line_no": 1,
                "article": "ART-001",
                "name": "Кабель ВВГнг",
                "unit": "м",
                "qty": 100,
                "nmc_unit_price": 150.0,
            },
            {
                "line_no": 2,
                "article": "ART-002",
                "name": "Кабель АВВГ",
                "unit": "м",
                "qty": 200,
                "nmc_unit_price": 250.0,
            },
        ],
    }
    calc = {
        "vat_rate": 20.0,
        "total_with_vat": 432600.0,
        "items": [
            {
                "line_no": 1,
                "quote_name": "Кабель ВВГнг",
                "unit_price_wo_vat": 120.0,
                "line_total_wo_vat": 12000.0,
            },
            {
                "line_no": 2,
                "quote_name": "Кабель АВВГ",
                "unit_price_wo_vat": 200.0,
                "line_total_wo_vat": 40000.0,
            },
        ],
    }

    for name, data in [
        ("company_profile.json", profile),
        ("tender.json", tender),
        ("calc.json", calc),
    ]:
        (data_dir / name).write_text(
            json.dumps(data, ensure_ascii=False), encoding="utf-8"
        )


def _create_mapping(
    mappings_dir: Path,
    template_path: str,
    output_name: str,
    fields: list[dict] | None = None,
    table_rows: list[dict] | None = None,
) -> None:
    """Создаёт YAML-маппинг в директории."""
    mapping = {
        "document": {
            "name": "Тестовый документ",
            "template": template_path,
            "output_name": output_name,
        },
        "fields": fields or [],
        "table_rows": table_rows or [],
    }
    (mappings_dir / "01_test.yaml").write_text(
        yaml.dump(mapping, allow_unicode=True), encoding="utf-8"
    )


# ---------------------------------------------------------------------------
# Тесты
# ---------------------------------------------------------------------------

class TestGenerateDocument:
    """Тесты для DocumentGenerator.generate_document."""

    def test_generate_document(self, tmp_path: Path) -> None:
        """Генерация документа - файл создан, плейсхолдеры заменены."""
        # Подготовка: данные, шаблон, маппинг
        data_dir = tmp_path / "data"
        data_dir.mkdir()
        _create_test_jsons(data_dir)

        templates_dir = tmp_path / "templates"
        templates_dir.mkdir()
        template_path = templates_dir / "test.docx"
        _create_test_docx(template_path)

        mappings_dir = tmp_path / "mappings"
        mappings_dir.mkdir()
        output_dir = tmp_path / "output"

        _create_mapping(
            mappings_dir,
            template_path=str(template_path),
            output_name="result.docx",
            fields=[
                {
                    "placeholder": "[NAME]",
                    "source": "PROFILE",
                    "path": "company.full_name",
                    "required": True,
                },
                {
                    "placeholder": "[TOTAL]",
                    "source": "CALC",
                    "path": "total_with_vat",
                    "transform": "money",
                },
            ],
        )

        gen = DocumentGenerator(
            data_dir=data_dir,
            mappings_dir=mappings_dir,
            output_dir=output_dir,
        )
        ctx = gen._load_data()
        mapping = gen._load_mappings()[0]
        result_path = gen.generate_document(mapping, ctx)

        # Файл создан
        assert result_path.exists()

        # Проверяем содержимое
        doc = Document(str(result_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "ООО «Тестовая Компания»" in full_text
        assert "432 600,00" in full_text
        assert "[NAME]" not in full_text

    def test_missing_required_field(self, tmp_path: Path) -> None:
        """Отсутствующее обязательное поле → маркер [НЕ УКАЗАНО] в документе."""
        data_dir = tmp_path / "data"
        data_dir.mkdir()
        _create_test_jsons(data_dir)

        templates_dir = tmp_path / "templates"
        templates_dir.mkdir()
        template_path = templates_dir / "test.docx"
        _create_test_docx(template_path)

        mappings_dir = tmp_path / "mappings"
        mappings_dir.mkdir()
        output_dir = tmp_path / "output"

        _create_mapping(
            mappings_dir,
            template_path=str(template_path),
            output_name="result.docx",
            fields=[
                {
                    "placeholder": "[NAME]",
                    "source": "PROFILE",
                    "path": "nonexistent.field.path",
                    "required": True,
                },
            ],
        )

        gen = DocumentGenerator(
            data_dir=data_dir,
            mappings_dir=mappings_dir,
            output_dir=output_dir,
        )
        ctx = gen._load_data()
        mapping = gen._load_mappings()[0]
        result_path = gen.generate_document(mapping, ctx)

        doc = Document(str(result_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[НЕ УКАЗАНО]" in full_text

    def test_transform_money(self, tmp_path: Path) -> None:
        """Трансформация money - число форматируется с пробелами и запятой."""
        data_dir = tmp_path / "data"
        data_dir.mkdir()
        _create_test_jsons(data_dir)

        templates_dir = tmp_path / "templates"
        templates_dir.mkdir()
        template_path = templates_dir / "test.docx"
        _create_test_docx(template_path)

        mappings_dir = tmp_path / "mappings"
        mappings_dir.mkdir()
        output_dir = tmp_path / "output"

        _create_mapping(
            mappings_dir,
            template_path=str(template_path),
            output_name="result.docx",
            fields=[
                {
                    "placeholder": "[TOTAL]",
                    "source": "CALC",
                    "path": "total_with_vat",
                    "transform": "money",
                },
            ],
        )

        gen = DocumentGenerator(
            data_dir=data_dir,
            mappings_dir=mappings_dir,
            output_dir=output_dir,
        )
        ctx = gen._load_data()
        mapping = gen._load_mappings()[0]
        result_path = gen.generate_document(mapping, ctx)

        doc = Document(str(result_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "432 600,00" in full_text


class TestGenerateAll:
    """Тесты для DocumentGenerator.generate_all."""

    def test_generate_all(self, tmp_path: Path) -> None:
        """generate_all - обрабатывает все маппинги, возвращает список путей."""
        data_dir = tmp_path / "data"
        data_dir.mkdir()
        _create_test_jsons(data_dir)

        templates_dir = tmp_path / "templates"
        templates_dir.mkdir()
        template_path = templates_dir / "test.docx"
        _create_test_docx(template_path)

        mappings_dir = tmp_path / "mappings"
        mappings_dir.mkdir()
        output_dir = tmp_path / "output"

        _create_mapping(
            mappings_dir,
            template_path=str(template_path),
            output_name="result.docx",
            fields=[
                {
                    "placeholder": "[NAME]",
                    "source": "PROFILE",
                    "path": "company.full_name",
                },
            ],
        )

        gen = DocumentGenerator(
            data_dir=data_dir,
            mappings_dir=mappings_dir,
            output_dir=output_dir,
        )
        paths = gen.generate_all()

        assert len(paths) == 1
        assert all(p.exists() for p in paths)
        assert paths[0].name == "result.docx"
