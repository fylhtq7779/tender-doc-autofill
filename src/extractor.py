"""Извлечение данных тендера из DOCX-документа ТКП (Запрос ТКП).

Детерминированный парсинг на основе regex - без использования LLM.
Результат: словарь, совместимый со структурой tender.json.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document


class TenderExtractor:
    """Извлекает структурированные данные тендера из DOCX-файла ТКП."""

    def __init__(self, docx_path: Path | str) -> None:
        self._path = Path(docx_path)
        if not self._path.exists():
            raise FileNotFoundError(f"Файл не найден: {self._path}")
        self._doc = Document(str(self._path))

    def extract(self) -> dict[str, Any]:
        """Извлекает все доступные поля из ТКП и возвращает словарь."""
        paragraphs = [p.text for p in self._doc.paragraphs]

        result: dict[str, Any] = {
            "purchase_number": self._extract_field(paragraphs, r"Номер закупки:\s*(\S+)"),
            "lot_number": self._extract_field(paragraphs, r"Лот:\s*(.+?)(?:\s{2,}|Код лота:|$)"),
            "lot_code": self._extract_field(paragraphs, r"Код лота:\s*(\S+)"),
            "subject": self._extract_field(paragraphs, r"Предмет закупки:\s*(.+)"),
            "bid_deadline": self._extract_field(paragraphs, r"Срок подачи предложений:\s*(.+)"),
            "offer_validity_days": None,
            "contract_number": None,
            "currency": "RUB",
            "customer": self._extract_customer(paragraphs),
            "delivery": self._extract_delivery(paragraphs),
            "payment": self._extract_payment(paragraphs),
            "warranty": self._extract_warranty(paragraphs),
            "items": self._extract_items(),
        }
        return result

    # ------------------------------------------------------------------
    # Приватные методы извлечения
    # ------------------------------------------------------------------

    def _extract_field(self, paragraphs: list[str], pattern: str) -> str:
        """Ищет первое совпадение regex среди всех параграфов."""
        for text in paragraphs:
            m = re.search(pattern, text)
            if m:
                return m.group(1).strip()
        return ""

    def _extract_customer(self, paragraphs: list[str]) -> dict[str, Any]:
        """Извлекает данные заказчика из строки 'Заказчик: ...'."""
        full_name = self._extract_field(paragraphs, r"Заказчик:\s*(.+)")
        return {
            "full_name": full_name,
        }

    def _extract_delivery(self, paragraphs: list[str]) -> dict[str, Any]:
        """Извлекает условия доставки."""
        place = self._extract_field(paragraphs, r"Место поставки:\s*(.+?)\.?\s*$")
        term_text = self._extract_field(paragraphs, r"Срок поставки:\s*(.+?)\.?\s*$")
        return {
            "place": place,
            "term_text": term_text,
        }

    def _extract_payment(self, paragraphs: list[str]) -> dict[str, Any]:
        """Извлекает условия оплаты."""
        term_text = self._extract_field(paragraphs, r"(Оплата\s+.+?)\.?\s*$")
        return {
            "term_text": term_text,
        }

    def _extract_warranty(self, paragraphs: list[str]) -> dict[str, Any]:
        """Извлекает гарантийные условия."""
        term_text = self._extract_field(paragraphs, r"(Гарантийный срок\s+.+?)\.?\s*$")
        return {
            "term_text": term_text,
        }

    def _extract_items(self) -> list[dict[str, Any]]:
        """Извлекает позиции из первой таблицы документа."""
        if not self._doc.tables:
            return []

        table = self._doc.tables[0]
        items: list[dict[str, Any]] = []

        # Пропускаем заголовок (row 0), парсим строки данных
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 8:
                continue

            # Пропускаем пустые строки
            if not cells[0]:
                continue

            item: dict[str, Any] = {
                "line_no": self._parse_int(cells[0]),
                "customer_name_code": cells[1],
                "article": cells[2],
                "name": cells[3],
                "qty": self._parse_int(cells[4]),
                "unit": cells[5],
                "nmc_unit_price": self._parse_float(cells[6]),
                "required_delivery_date": cells[7],
            }
            items.append(item)

        return items

    @staticmethod
    def _parse_float(value: str) -> float:
        """Парсит число с запятой как десятичным разделителем: '185,00' -> 185.0."""
        cleaned = value.replace("\xa0", "").replace(" ", "").replace(",", ".")
        try:
            return float(cleaned)
        except ValueError:
            return 0.0

    @staticmethod
    def _parse_int(value: str) -> int:
        """Парсит целое число из строки: '1200' -> 1200."""
        cleaned = value.replace("\xa0", "").replace(" ", "")
        try:
            return int(cleaned)
        except ValueError:
            return 0
