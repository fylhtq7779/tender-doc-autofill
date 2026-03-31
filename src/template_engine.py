"""Движок подстановки значений в DOCX-шаблоны.

Класс TemplateEngine загружает DOCX через python-docx и предоставляет методы:
  - replace_placeholder — замена плейсхолдера на значение во всём документе
  - fill_table_row      — заполнение строки таблицы по индексам колонок
  - save                — сохранение результата в файл
"""

from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


class TemplateEngine:
    """Работа с DOCX-шаблоном: поиск и замена плейсхолдеров, заполнение таблиц."""

    def __init__(self, template_path: Path) -> None:
        """Загружает копию шаблона в память.

        Args:
            template_path: путь к DOCX-шаблону
        """
        self._template_path = template_path
        # Загружаем документ — python-docx работает с копией в памяти,
        # но для гарантии неизменности оригинала перечитываем из файла
        self._doc: Document = Document(str(template_path))

    # ------------------------------------------------------------------
    # Публичные методы
    # ------------------------------------------------------------------

    def replace_placeholder(self, placeholder: str, value: str) -> int:
        """Заменяет плейсхолдер на значение во всём документе.

        Ищет в параграфах верхнего уровня и во всех ячейках таблиц.
        Замена происходит на уровне Run — форматирование сохраняется.

        Args:
            placeholder: строка для поиска, например "[COMPANY]"
            value: строка для подстановки

        Returns:
            количество произведённых замен
        """
        count = 0

        # 1. Параграфы верхнего уровня документа
        for paragraph in self._doc.paragraphs:
            count += self._replace_in_paragraph(paragraph, placeholder, value)

        # 2. Все таблицы → строки → ячейки → параграфы
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        count += self._replace_in_paragraph(
                            paragraph, placeholder, value
                        )

        return count

    def fill_table_row(
        self,
        table_idx: int,
        row_idx: int,
        col_values: dict[int, str],
    ) -> None:
        """Заполняет ячейки строки таблицы по индексам колонок.

        Форматирование первого Run в ячейке сохраняется.

        Args:
            table_idx: индекс таблицы в документе
            row_idx: индекс строки в таблице
            col_values: словарь {индекс_колонки: значение}

        Raises:
            IndexError: если table_idx или row_idx выходят за пределы
        """
        table = self._doc.tables[table_idx]
        row = table.rows[row_idx]

        for col_idx, value in col_values.items():
            cell = row.cells[col_idx]
            self._set_cell_text(cell, value)

    def save(self, output_path: Path) -> None:
        """Сохраняет документ в файл. Создаёт директории при необходимости.

        Args:
            output_path: путь для сохранения результата
        """
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(str(output_path))

    # ------------------------------------------------------------------
    # Приватные методы
    # ------------------------------------------------------------------

    def _replace_in_paragraph(
        self, paragraph: Paragraph, placeholder: str, value: str
    ) -> int:
        """Заменяет плейсхолдер в Run-ах параграфа. Возвращает число замен."""
        count = 0
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)
                count += 1
        return count

    @staticmethod
    def _set_cell_text(cell: object, value: str) -> None:
        """Устанавливает текст ячейки, сохраняя форматирование первого Run."""
        # Берём первый параграф ячейки
        paragraph = cell.paragraphs[0]
        runs = paragraph.runs

        if runs:
            # Сохраняем первый Run (с его форматированием), остальные удаляем
            first_run = runs[0]
            first_run.text = value

            # Удаляем лишние Run-ы из XML
            for run in runs[1:]:
                run._element.getparent().remove(run._element)
        else:
            # Нет Run-ов — добавляем текст напрямую
            paragraph.add_run(value)

        # Удаляем лишние параграфы (ячейка может содержать несколько)
        for p in cell.paragraphs[1:]:
            p._element.getparent().remove(p._element)
